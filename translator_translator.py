pip install requests python-docx

# Imports
import os
import uuid
import requests
from urllib.parse import urljoin
from docx import Document

# Configs
SUBSCRIPTION_KEY = 'My_Key'
ENDPOINT = 'My_Endpoint'
LOCATION = 'eastus'
API_VERSION = '3.0'

### Parte 1

def translate_text(text, source_language='en', target_language='pt-br'):
  """
  Traduz o texto fornecido usando a API do Azure Translator.

  Args:
    text (str): O texto a ser traduzido.
    source_language (str): Código do idioma de origem (padrão: 'en').
    target_language (str): Código do idioma de destino (padrão: 'pt-br').

  Returns:
    str: Texto traduzido.
  """
    
  if not text:
    raise ValueError("O texto a ser traduzido não pode ser vazio.")
  if not target_language:
    raise ValueError("A linguagem de destino deve ser especificada.")
    
  path = '/translate'
  constructed_url = urljoin(ENDPOINT, path)
    
  headers = {
    'Ocp-Apim-Subscription-key': SUBSCRIPTION_KEY,
    'Ocp-Apim-Subscription-Region': LOCATION,
    'Content-type': 'application/json',
    'X-ClientTraceId': str(uuid.uuid4())
  }

  body = [{'text': text}]
  params = {
    'api-version': API_VERSION,
    'from': source_language, # Define o idioma de origem
    'to': target_language # Define o idioma de destino
  }

  try:
    response = requests.post(constructed_url, params=params, headers=headers, json=body)
    response.raise_for_status()  # Garante que um erro HTTP seja levantado
    result = response.json()
    return result[0]['translations'][0]['text']
  except requests.exceptions.RequestException as e:
    raise RuntimeError(f"Erro na requisição: {e}")
  except (KeyError, IndexError) as e:
    raise RuntimeError(f"Erro ao processar a resposta da API: {e}")

### Parte 2

def translate_doc(path, target_language='pt-br', source_language='en'):
  """
  Traduz um documento do Word (.docx) usando a função de tradução.

  Args:
    path (str): Caminho para o arquivo .docx de entrada.
    target_language (str): Código do idioma de destino (padrão: 'pt-br').
    source_language (str): Código do idioma de origem (padrão: 'en').

  Returns:
    str: Caminho para o arquivo traduzido salvo.
  """
  if not os.path.exists(path):
    raise FileNotFoundError(f"O arquivo '{path}' não foi encontrado.")
  if not path.lower().endswith('.docx'):
    raise ValueError("O arquivo fornecido deve estar no formato .docx.")
    
  # Lê o documento de entrada
  document = Document(path)
  full_text = []

  # Processa parágrafo por parágrafo
  for paragraph in document.paragraphs:
    if paragraph.text.strip():  # Ignora parágrafos vazios
      translated_text = translate_text(paragraph.text, target_language=target_language, source_language=source_language)
      full_text.append(translated_text)
    else:
      full_text.append("")  # Mantém a estrutura do documento

  # Cria um novo documento traduzido
  translated_doc = Document()
  for line in full_text:
    translated_doc.add_paragraph(line)

  # Define o caminho para o arquivo traduzido
  path_translated = path.replace('.docx', f'_{target_language}.docx')
  translated_doc.save(path_translated)
    
  return path_translated
